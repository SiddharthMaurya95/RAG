import os
import sqlite3
import datetime
import re
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.paths import get_db_path

class NumberedCanvas(canvas.Canvas):
    """Custom canvas that computes total page count dynamically and draws headers/footers."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_decorations(self, page_count):
        self.saveState()
        
        # Primary Top Accent Bar
        self.setFillColor(colors.HexColor('#1e3a8a')) # Deep Navy
        self.rect(54, 742, 504, 3, fill=True, stroke=False)
        
        # Header Text
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(colors.HexColor('#475569'))
        self.drawString(54, 752, "AUTOMOTIVE QUALITY INTELLIGENCE")
        
        self.setFont("Helvetica", 8)
        self.drawRightString(558, 752, datetime.datetime.now().strftime('%Y-%m-%d'))
        
        # Footer Divider Line
        self.setStrokeColor(colors.HexColor('#e2e8f0'))
        self.setLineWidth(0.5)
        self.line(54, 50, 558, 50)
        
        # Footer Text & Page Number
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor('#64748b'))
        self.drawString(54, 35, "CONFIDENTIAL - FOR INTERNAL USE ONLY")
        self.drawRightString(558, 35, f"Page {self._pageNumber} of {page_count}")
        
        self.restoreState()


class ReportEngine:
    def __init__(self, db_path="data/automotive.db"):
        self.db_path = get_db_path(db_path)

    def _get_report_data(self, year, month):
        """Retrieves statistics and summary cases for a given month and year."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # 1. Total claims count
        cursor.execute("SELECT COUNT(*) FROM records WHERE report_year = ? AND report_month = ?;", (year, month))
        total_claims = cursor.fetchone()[0]
        
        # 2. Claims by country
        cursor.execute("""
            SELECT outbreak_country, COUNT(*) as cnt
            FROM records
            WHERE report_year = ? AND report_month = ?
            GROUP BY outbreak_country
            ORDER BY cnt DESC
            LIMIT 15;
        """, (year, month))
        countries_data = cursor.fetchall()
        
        # 3. Claims by model
        cursor.execute("""
            SELECT product_model_code, COUNT(*) as cnt
            FROM records
            WHERE report_year = ? AND report_month = ?
            GROUP BY product_model_code
            ORDER BY cnt DESC
            LIMIT 15;
        """, (year, month))
        models_data = cursor.fetchall()
        
        # 4. Top trouble codes
        cursor.execute("""
            SELECT trouble_code_complaint, COUNT(*) as cnt
            FROM records
            WHERE report_year = ? AND report_month = ? AND trouble_code_complaint IS NOT NULL
            GROUP BY trouble_code_complaint
            ORDER BY cnt DESC
            LIMIT 5;
        """, (year, month))
        trouble_codes_data = cursor.fetchall()
        
        # 5. Top 5 detailed cases with summaries
        cursor.execute("""
            SELECT ftir_no, subject, quality, reported_company, summary
            FROM records
            WHERE report_year = ? AND report_month = ?
            LIMIT 5;
        """, (year, month))
        cases_data = cursor.fetchall()
        
        conn.close()
        
        return {
            "total_claims": total_claims,
            "countries": countries_data,
            "models": models_data,
            "trouble_codes": trouble_codes_data,
            "cases": cases_data
        }

    def _markdown_to_html(self, text):
        """Helper to convert basic Markdown to ReportLab HTML-like tags."""
        if not text:
            return ""
        text = html_escape(text)
        # Bold **text** -> <b>text</b>
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        # Italic *text* -> <i>text</i>
        text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
        # Code backticks `code` -> <font name="Courier">code</font>
        text = re.sub(r'`(.*?)`', r'<font name="Courier" color="#c7254e" size="9"><b>\1</b></font>', text)
        # Newlines
        text = text.replace('\n', '<br/>')
        return text

    def generate_pdf_report(self, year, month, output_path):
        """Generates a professional PDF report using ReportLab with premium styling."""
        data = self._get_report_data(year, month)
        month_name = datetime.date(1900, month, 1).strftime('%B')
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54, leftMargin=54,
            topMargin=72, bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        # Premium Styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=24,
            textColor=colors.HexColor('#0f172a'), # Slate 900
            spaceAfter=6,
            keepWithNext=True
        )
        
        subtitle_style = ParagraphStyle(
            'ReportSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            textColor=colors.HexColor('#475569'), # Slate 600
            spaceAfter=15,
            keepWithNext=True
        )
        
        h2_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            textColor=colors.HexColor('#1e3a8a'), # Navy
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True
        )
        
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            textColor=colors.HexColor('#334155'), # Slate 700
            leading=15,
            spaceAfter=8
        )
        
        table_hdr_style = ParagraphStyle(
            'TableHdr',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=9,
            textColor=colors.HexColor('#ffffff'),
            leading=12
        )
        
        table_text_style = ParagraphStyle(
            'TableText',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            textColor=colors.HexColor('#334155'),
            leading=13
        )
        
        story = []
        
        # Header / Title Block
        story.append(Paragraph(f"Monthly QA Intelligence Report", title_style))
        story.append(Paragraph(f"Reporting Period: <b>{month_name} {year}</b> | Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style))
        story.append(Spacer(1, 10))
        
        # Executive Summary Box
        story.append(Paragraph("Executive Summary", h2_style))
        summary_text = (
            f"During the period of {month_name} {year}, a total of <b>{data['total_claims']}</b> failure claims "
            f"were logged in the system. The breakdown of models, outbreaks by country, and critical trouble codes "
            f"indicates the current quality distribution across manufacturing bases."
        )
        
        summary_table = Table([[Paragraph(summary_text, body_style)]], colWidths=[504])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
            ('PADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 10))
        
        # Data Breakdown tables side-by-side
        story.append(Paragraph("Quality Distribution Metrics", h2_style))
        
        grid_data = [
            [Paragraph("Outbreaks by Country", table_hdr_style), Paragraph("Claims by Product Model", table_hdr_style)]
        ]
        
        country_lines = "<br/>".join([f"• {c[0]}: <b>{c[1]}</b> claims" for c in data['countries']]) or "No data available."
        model_lines = "<br/>".join([f"• {m[0]}: <b>{m[1]}</b> claims" for m in data['models']]) or "No data available."
        
        grid_data.append([Paragraph(country_lines, table_text_style), Paragraph(model_lines, table_text_style)])
        
        grid_table = Table(grid_data, colWidths=[250, 254])
        grid_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,0), 6),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white]),
            ('PADDING', (0,0), (-1,-1), 10)
        ]))
        story.append(grid_table)
        story.append(Spacer(1, 12))
        
        # Trouble Codes
        story.append(Paragraph("Top Failure Trouble Codes (DTC)", h2_style))
        tc_rows = [[Paragraph("Trouble Code", table_hdr_style), Paragraph("Claims Count", table_hdr_style)]]
        for row in data['trouble_codes']:
            tc_rows.append([Paragraph(str(row[0]), table_text_style), Paragraph(str(row[1]), table_text_style)])
            
        if len(tc_rows) == 1:
            tc_rows.append([Paragraph("None", table_text_style), Paragraph("0", table_text_style)])
            
        tc_table = Table(tc_rows, colWidths=[354, 150])
        tc_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#475569')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
            ('PADDING', (0,0), (-1,-1), 6)
        ]))
        story.append(tc_table)
        story.append(Spacer(1, 15))
        
        # Featured Quality Cases
        story.append(Paragraph("Sample Quality Cases & Technical Summaries", h2_style))
        for r in data['cases']:
            ftir_no, subject, quality, dealer, summary = r
            case_header = f"<b>FTIR No: {ftir_no}</b> | Dealer: {dealer} | Quality rating: <b>{quality}</b>"
            story.append(Paragraph(case_header, body_style))
            story.append(Paragraph(f"<i>Subject:</i> {subject}", table_text_style))
            story.append(Paragraph(f"<i>Summary:</i> {summary}", body_style))
            story.append(Spacer(1, 8))
            
        if not data['cases']:
            story.append(Paragraph("No detailed cases available for this month.", body_style))
            
        doc.build(story, canvasmaker=NumberedCanvas)
        print(f"PDF report generated at: {output_path}")

    def generate_docx_report(self, year, month, output_path):
        """Generates a professional DOCX report using python-docx with custom styles."""
        data = self._get_report_data(year, month)
        month_name = datetime.date(1900, month, 1).strftime('%B')
        
        doc = Document()
        
        # Document Title
        title = doc.add_paragraph()
        run_title = title.add_run('Automotive QA Intelligence Monthly Report')
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(20)
        run_title.bold = True
        run_title.font.color.rgb = colors.HexColor('#0f172a') # Slate 900
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p = doc.add_paragraph()
        p.add_run(f"Reporting Period: {month_name} {year}\n").bold = True
        p.add_run(f"Generated at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Executive Summary
        h1 = doc.add_heading(level=1)
        r1 = h1.add_run('Executive Summary')
        r1.font.name = 'Arial'
        r1.font.color.rgb = colors.HexColor('#1e3a8a')
        
        doc.add_paragraph(
            f"During the period of {month_name} {year}, a total of {data['total_claims']} failure claims "
            f"were logged in the system. The breakdown of models, outbreaks by country, and critical trouble codes "
            f"indicates the current quality distribution across manufacturing bases."
        )
        
        # Outbreaks by Country
        h2 = doc.add_heading(level=2)
        r2 = h2.add_run('Outbreaks by Country')
        r2.font.name = 'Arial'
        
        table_country = doc.add_table(rows=1, cols=2)
        table_country.style = 'Light Shading Accent 1'
        hdr_cells = table_country.rows[0].cells
        hdr_cells[0].text = 'Country'
        hdr_cells[1].text = 'Claims Count'
        for country, count in data['countries']:
            row_cells = table_country.add_row().cells
            row_cells[0].text = str(country)
            row_cells[1].text = str(count)
            
        # Claims by Model Code
        h3 = doc.add_heading(level=2)
        r3 = h3.add_run('Claims by Product Model')
        r3.font.name = 'Arial'
        
        table_model = doc.add_table(rows=1, cols=2)
        table_model.style = 'Light Shading Accent 1'
        hdr_cells_m = table_model.rows[0].cells
        hdr_cells_m[0].text = 'Model Code'
        hdr_cells_m[1].text = 'Claims Count'
        for model, count in data['models']:
            row_cells = table_model.add_row().cells
            row_cells[0].text = str(model)
            row_cells[1].text = str(count)
            
        # Top Trouble Codes
        h4 = doc.add_heading(level=2)
        r4 = h4.add_run('Top Failure Trouble Codes (DTC)')
        r4.font.name = 'Arial'
        
        table_tc = doc.add_table(rows=1, cols=2)
        table_tc.style = 'Light Shading Accent 1'
        hdr_cells_tc = table_tc.rows[0].cells
        hdr_cells_tc[0].text = 'Trouble Code'
        hdr_cells_tc[1].text = 'Claims Count'
        for tc, count in data['trouble_codes']:
            row_cells = table_tc.add_row().cells
            row_cells[0].text = str(tc)
            row_cells[1].text = str(count)
            
        # Featured Quality Cases
        h5 = doc.add_heading(level=1)
        r5 = h5.add_run('Sample Quality Cases & Technical Summaries')
        r5.font.name = 'Arial'
        r5.font.color.rgb = colors.HexColor('#1e3a8a')
        
        for r in data['cases']:
            ftir_no, subject, quality, dealer, summary = r
            p_case = doc.add_paragraph()
            p_case.add_run(f"FTIR No: {ftir_no}\n").bold = True
            p_case.add_run(f"Dealer: {dealer} | Quality rating: {quality}\n")
            p_case.add_run(f"Subject: {subject}\n").italic = True
            p_case.add_run(f"Summary: {summary}\n")
            
        if not data['cases']:
            doc.add_paragraph("No detailed cases available for this month.")
            
        doc.save(output_path)
        print(f"DOCX report generated at: {output_path}")

    def generate_chat_pdf(self, session_title, chat_history, username, output_path):
        """Generates a beautifully formatted PDF transcript of the active chat session."""
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54, leftMargin=54,
            topMargin=72, bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'ChatTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=4,
            keepWithNext=True
        )
        
        meta_style = ParagraphStyle(
            'ChatMeta',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            textColor=colors.HexColor('#64748b'),
            spaceAfter=15,
            keepWithNext=True
        )
        
        user_header_style = ParagraphStyle(
            'UserHeader',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10,
            textColor=colors.HexColor('#1e3a8a'),
            spaceBefore=12,
            spaceAfter=2,
            keepWithNext=True
        )
        
        user_body_style = ParagraphStyle(
            'UserBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            textColor=colors.HexColor('#1e293b'),
            leading=14
        )
        
        assistant_header_style = ParagraphStyle(
            'AssistantHeader',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10,
            textColor=colors.HexColor('#0d9488'),
            spaceBefore=12,
            spaceAfter=2,
            keepWithNext=True
        )
        
        assistant_body_style = ParagraphStyle(
            'AssistantBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            textColor=colors.HexColor('#334155'),
            leading=14,
            spaceAfter=6
        )
        
        code_style = ParagraphStyle(
            'CodeStyle',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=8.5,
            textColor=colors.HexColor('#38bdf8'),
            leading=11
        )
        
        citation_header_style = ParagraphStyle(
            'CitationHeader',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=8.5,
            textColor=colors.HexColor('#0284c7'),
            spaceBefore=4,
            spaceAfter=2,
            keepWithNext=True
        )
        
        citation_body_style = ParagraphStyle(
            'CitationBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=8,
            textColor=colors.HexColor('#475569'),
            leading=11
        )
        
        story = []
        
        # Document Title Block
        story.append(Paragraph(f"Chat Session Transcript", title_style))
        story.append(Paragraph(f"Session: <b>{session_title}</b> | Operator: {username} | Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", meta_style))
        story.append(Spacer(1, 10))
        
        for i, msg in enumerate(chat_history):
            role = msg["role"]
            content_html = self._markdown_to_html(msg["content"])
            
            if role == "user":
                # User Bubble Table Box
                user_content_p = Paragraph(content_html, user_body_style)
                user_box = Table([[user_content_p]], colWidths=[504])
                user_box.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#eff6ff')), # Light blue bubble
                    ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#bfdbfe')),
                    ('PADDING', (0,0), (-1,-1), 8),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ]))
                story.append(Paragraph(f"👤 Operator:", user_header_style))
                story.append(user_box)
                story.append(Spacer(1, 8))
                
            else:
                # Assistant message
                intent = msg.get("intent", "GENERAL")
                story.append(Paragraph(f"✨ Assistant | <font size='8'>[{intent}]</font>", assistant_header_style))
                story.append(Paragraph(content_html, assistant_body_style))
                
                # SQL Query
                sql_query = msg.get("sql_query")
                if sql_query:
                    sql_html = html_escape(sql_query).replace('\n', '<br/>')
                    sql_p = Paragraph(sql_html, code_style)
                    sql_box = Table([[sql_p]], colWidths=[504])
                    sql_box.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#0f172a')), # Dark block
                        ('PADDING', (0,0), (-1,-1), 8),
                        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#1e293b')),
                    ]))
                    story.append(Paragraph("<b>Executed SQL Query:</b>", assistant_body_style))
                    story.append(sql_box)
                    story.append(Spacer(1, 6))
                
                # Tabular DataFrame Data
                df = msg.get("df")
                if df is not None and not df.empty:
                    story.append(Paragraph("<b>Query Results Table:</b>", assistant_body_style))
                    
                    headers = [Paragraph(f"<b>{col}</b>", table_hdr_style_from_pdf()) for col in df.columns]
                    data = [headers]
                    for _, row in df.iterrows():
                        data.append([Paragraph(html_escape(str(val)), table_text_style_from_pdf()) for val in row])
                    
                    col_widths = [504 / len(df.columns)] * len(df.columns)
                    df_table = Table(data, colWidths=col_widths)
                    df_table.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0d9488')), # Teal header
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
                        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
                        ('PADDING', (0,0), (-1,-1), 5),
                    ]))
                    story.append(df_table)
                    story.append(Spacer(1, 8))
                    
                # Chart Visuals
                chart_type = msg.get("chart_type")
                chart_title = msg.get("chart_title", "Chart")
                if chart_type and chart_type != "empty" and df is not None and not df.empty:
                    from reportlab.platypus import Image as RLImage
                    import tempfile
                    import time
                    
                    from viz.charts import (
                        plot_horizontal_bar,
                        plot_line_trend,
                        plot_donut_chart,
                        plot_histogram,
                        plot_radar_comparison,
                        plot_grouped_bar
                    )
                    
                    fig = None
                    try:
                        if chart_type == "horizontal_bar":
                            fig = plot_horizontal_bar(df, df.columns[1], df.columns[0], chart_title)
                        elif chart_type == "line":
                            y_cols = list(df.columns[1:])
                            fig = plot_line_trend(df, df.columns[0], y_cols, chart_title)
                        elif chart_type == "donut":
                            fig = plot_donut_chart(df, df.columns[0], df.columns[1], chart_title)
                        elif chart_type == "histogram":
                            fig = plot_histogram(df, df.columns[0], chart_title)
                        elif chart_type == "radar":
                            metrics = list(df.columns[1:])
                            fig = plot_radar_comparison(df, df.columns[0], metrics, chart_title)
                        elif chart_type == "grouped_bar":
                            metrics = list(df.columns[1:])
                            fig = plot_grouped_bar(df, df.columns[0], metrics, chart_title)
                        
                        if fig:
                            temp_img_path = os.path.join(tempfile.gettempdir(), f"chart_{i}_{int(time.time())}.png")
                            fig.write_image(temp_img_path, format="png", width=700, height=350)
                            story.append(Paragraph(f"<b>Chart Visualization: {chart_title}</b>", assistant_body_style))
                            story.append(RLImage(temp_img_path, width=450, height=225))
                            story.append(Spacer(1, 10))
                    except Exception as e:
                        print(f"Error rendering chart to PDF: {e}")
                
                # Citations
                citations = msg.get("citations")
                if citations:
                    story.append(Paragraph("Citations & Sources:", citation_header_style))
                    for cite in citations:
                        cite_text = f"• <b>FTIR {cite.get('ftir_no')}</b> | Subject: {cite.get('subject')} | Dealer: {cite.get('reported_company')} ({cite.get('outbreak_country')})"
                        story.append(Paragraph(cite_text, citation_body_style))
                    story.append(Spacer(1, 6))
                
                story.append(Spacer(1, 10))
                
        doc.build(story, canvasmaker=NumberedCanvas)
        print(f"Chat PDF transcript generated at: {output_path}")


def html_escape(text):
    """Escapes HTML special characters for ReportLab Paragraph compatibility."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def table_hdr_style_from_pdf():
    styles = getSampleStyleSheet()
    return ParagraphStyle(
        'TblHdr',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        textColor=colors.white,
        leading=11
    )

def table_text_style_from_pdf():
    styles = getSampleStyleSheet()
    return ParagraphStyle(
        'TblTxt',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        textColor=colors.HexColor('#334155'),
        leading=11
    )
