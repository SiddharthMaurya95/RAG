# =====================================================
# ✅ REPORT ENGINE
# =====================================================
import os
import sqlite3
import datetime
import re
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.paths import get_db_path
from core.utils.decorators import with_logging_and_exceptions

class NumberedCanvas(canvas.Canvas):
    """Custom canvas that computes total page count dynamically and draws headers/footers."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_decorations(self, page_count):
        self.saveState()
        
        # Primary Top Accent Bar
        self.setFillColor(colors.HexColor('#1e3a8a')) # Deep Navy
        self.rect(54, 742, 504, 3, fill=True, stroke=False)
        
        # Header Text
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(colors.HexColor('#475569'))
        self.drawString(54, 752, "AUTOMOTIVE QUALITY INTELLIGENCE")
        
        self.setFont("Helvetica", 8)
        self.drawRightString(558, 752, datetime.datetime.now().strftime('%Y-%m-%d'))
        
        # Footer Divider Line
        self.setStrokeColor(colors.HexColor('#e2e8f0'))
        self.setLineWidth(0.5)
        self.line(54, 50, 558, 50)
        
        # Footer Text & Page Number
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor('#64748b'))
        self.drawString(54, 35, "CONFIDENTIAL - FOR INTERNAL USE ONLY")
        self.drawRightString(558, 35, f"Page {self._pageNumber} of {page_count}")
        
        self.restoreState()


class ReportEngine:
    def __init__(self, db_path="data/automotive.db"):
        self.db_path = get_db_path(db_path)

    def _get_report_data(self, year, month):
        """Retrieves statistics and summary cases for a given month and year."""
        from core.database import get_engine
        conn = get_engine(self.db_path).raw_connection()
        cursor = conn.cursor()

        
        # 1. Total claims count
        cursor.execute("SELECT COUNT(*) FROM records WHERE report_year = ? AND report_month = ?;", (year, month))
        total_claims = cursor.fetchone()[0]
        
        # 2. Claims by country
        cursor.execute("""
            SELECT outbreak_country, COUNT(*) as cnt
            FROM records
            WHERE report_year = ? AND report_month = ?
            GROUP BY outbreak_country
            ORDER BY cnt DESC
            LIMIT 15;
        """, (year, month))
        countries_data = cursor.fetchall()
        
        # 3. Claims by model
        cursor.execute("""
            SELECT product_model_code, COUNT(*) as cnt
            FROM records
            WHERE report_year = ? AND report_month = ?
            GROUP BY product_model_code
            ORDER BY cnt DESC
            LIMIT 15;
        """, (year, month))
        models_data = cursor.fetchall()
        
        # 4. Top trouble codes
        cursor.execute("""
            SELECT trouble_code_complaint, COUNT(*) as cnt
            FROM records
            WHERE report_year = ? AND report_month = ? AND trouble_code_complaint IS NOT NULL
            GROUP BY trouble_code_complaint
            ORDER BY cnt DESC
            LIMIT 5;
        """, (year, month))
        trouble_codes_data = cursor.fetchall()
        
        # 5. Top 5 detailed cases with summaries
        cursor.execute("""
            SELECT ftir_no, subject, quality, reported_company, summary
            FROM records
            WHERE report_year = ? AND report_month = ?
            LIMIT 5;
        """, (year, month))
        cases_data = cursor.fetchall()
        
        conn.close()
        
        return {
            "total_claims": total_claims,
            "countries": countries_data,
            "models": models_data,
            "trouble_codes": trouble_codes_data,
            "cases": cases_data
        }

    def _markdown_to_html(self, text):
        """Helper to convert basic Markdown to ReportLab HTML-like tags."""
        if not text:
            return ""
        text = html_escape(text)
        # Bold **text** -> <b>text</b>
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        # Italic *text* -> <i>text</i>
        text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
        # Code backticks `code` -> <font name="Courier">code</font>
        text = re.sub(r'`(.*?)`', r'<font name="Courier" color="#c7254e" size="9"><b>\1</b></font>', text)
        # Newlines
        text = text.replace('\n', '<br/>')
        return text

    @with_logging_and_exceptions
    def generate_pdf_report(self, year, month, output_path):
        """Generates a professional PDF report using ReportLab with premium styling."""
        data = self._get_report_data(year, month)
        month_name = datetime.date(1900, month, 1).strftime('%B')
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54, leftMargin=54,
            topMargin=72, bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        # Premium Styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=24,
            textColor=colors.HexColor('#0f172a'), # Slate 900
            spaceAfter=6,
            keepWithNext=True
        )
        
        subtitle_style = ParagraphStyle(
            'ReportSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            textColor=colors.HexColor('#475569'), # Slate 600
            spaceAfter=15,
            keepWithNext=True
        )
        
        h2_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            textColor=colors.HexColor('#1e3a8a'), # Navy
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True
        )
        
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            textColor=colors.HexColor('#334155'), # Slate 700
            leading=15,
            spaceAfter=8
        )
        
        table_hdr_style = ParagraphStyle(
            'TableHdr',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=9,
            textColor=colors.HexColor('#ffffff'),
            leading=12
        )
        
        table_text_style = ParagraphStyle(
            'TableText',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            textColor=colors.HexColor('#334155'),
            leading=13
        )
        
        story = []
        
        # Header / Title Block
        story.append(Paragraph(f"Monthly QA Intelligence Report", title_style))
        story.append(Paragraph(f"Reporting Period: <b>{month_name} {year}</b> | Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style))
        story.append(Spacer(1, 10))
        
        # Executive Summary Box
        story.append(Paragraph("Executive Summary", h2_style))
        summary_text = (
            f"During the period of {month_name} {year}, a total of <b>{data['total_claims']}</b> failure claims "
            f"were logged in the system. The breakdown of models, outbreaks by country, and critical trouble codes "
            f"indicates the current quality distribution across manufacturing bases."
        )
        
        summary_table = Table([[Paragraph(summary_text, body_style)]], colWidths=[504])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
            ('PADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 10))
        
        # Data Breakdown tables side-by-side
        story.append(Paragraph("Quality Distribution Metrics", h2_style))
        
        grid_data = [
            [Paragraph("Outbreaks by Country", table_hdr_style), Paragraph("Claims by Product Model", table_hdr_style)]
        ]
        
        country_lines = "<br/>".join([f"• {c[0]}: <b>{c[1]}</b> claims" for c in data['countries']]) or "No data available."
        model_lines = "<br/>".join([f"• {m[0]}: <b>{m[1]}</b> claims" for m in data['models']]) or "No data available."
        
        grid_data.append([Paragraph(country_lines, table_text_style), Paragraph(model_lines, table_text_style)])
        
        grid_table = Table(grid_data, colWidths=[250, 254])
        grid_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,0), 6),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white]),
            ('PADDING', (0,0), (-1,-1), 10)
        ]))
        story.append(grid_table)
        story.append(Spacer(1, 12))
        
        # Trouble Codes
        story.append(Paragraph("Top Failure Trouble Codes (DTC)", h2_style))
        tc_rows = [[Paragraph("Trouble Code", table_hdr_style), Paragraph("Claims Count", table_hdr_style)]]
        for row in data['trouble_codes']:
            tc_rows.append([Paragraph(str(row[0]), table_text_style), Paragraph(str(row[1]), table_text_style)])
            
        if len(tc_rows) == 1:
            tc_rows.append([Paragraph("None", table_text_style), Paragraph("0", table_text_style)])
            
        tc_table = Table(tc_rows, colWidths=[354, 150])
        tc_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#475569')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
            ('PADDING', (0,0), (-1,-1), 6)
        ]))
        story.append(tc_table)
        story.append(Spacer(1, 15))
        
        # Featured Quality Cases
        story.append(Paragraph("Sample Quality Cases & Technical Summaries", h2_style))
        for r in data['cases']:
            ftir_no, subject, quality, dealer, summary = r
            case_header = f"<b>FTIR No: {ftir_no}</b> | Dealer: {dealer} | Quality rating: <b>{quality}</b>"
            story.append(Paragraph(case_header, body_style))
            story.append(Paragraph(f"<i>Subject:</i> {subject}", table_text_style))
            story.append(Paragraph(f"<i>Summary:</i> {summary}", body_style))
            story.append(Spacer(1, 8))
            
        if not data['cases']:
            story.append(Paragraph("No detailed cases available for this month.", body_style))
            
        doc.build(story, canvasmaker=NumberedCanvas)
        print(f"PDF report generated at: {output_path}")

    @with_logging_and_exceptions
    def generate_docx_report(self, year, month, output_path):
        """Generates a professional DOCX report using python-docx with custom styles."""
        data = self._get_report_data(year, month)
        month_name = datetime.date(1900, month, 1).strftime('%B')
        
        doc = Document()
        
        # Document Title
        title = doc.add_paragraph()
        run_title = title.add_run('Automotive QA Intelligence Monthly Report')
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(20)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p = doc.add_paragraph()
        p.add_run(f"Reporting Period: {month_name} {year}\n").bold = True
        p.add_run(f"Generated at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Executive Summary
        h1 = doc.add_heading(level=1)
        r1 = h1.add_run('Executive Summary')
        r1.font.name = 'Arial'
        r1.font.color.rgb = RGBColor(30, 58, 138)
        
        doc.add_paragraph(
            f"During the period of {month_name} {year}, a total of {data['total_claims']} failure claims "
            f"were logged in the system. The breakdown of models, outbreaks by country, and critical trouble codes "
            f"indicates the current quality distribution across manufacturing bases."
        )
        
        # Outbreaks by Country
        h2 = doc.add_heading(level=2)
        r2 = h2.add_run('Outbreaks by Country')
        r2.font.name = 'Arial'
        
        table_country = doc.add_table(rows=1, cols=2)
        table_country.style = 'Light Shading Accent 1'
        hdr_cells = table_country.rows[0].cells
        hdr_cells[0].text = 'Country'
        hdr_cells[1].text = 'Claims Count'
        for country, count in data['countries']:
            row_cells = table_country.add_row().cells
            row_cells[0].text = str(country)
            row_cells[1].text = str(count)
            
        # Claims by Model Code
        h3 = doc.add_heading(level=2)
        r3 = h3.add_run('Claims by Product Model')
        r3.font.name = 'Arial'
        
        table_model = doc.add_table(rows=1, cols=2)
        table_model.style = 'Light Shading Accent 1'
        hdr_cells_m = table_model.rows[0].cells
        hdr_cells_m[0].text = 'Model Code'
        hdr_cells_m[1].text = 'Claims Count'
        for model, count in data['models']:
            row_cells = table_model.add_row().cells
            row_cells[0].text = str(model)
            row_cells[1].text = str(count)
            
        # Top Trouble Codes
        h4 = doc.add_heading(level=2)
        r4 = h4.add_run('Top Failure Trouble Codes (DTC)')
        r4.font.name = 'Arial'
        
        table_tc = doc.add_table(rows=1, cols=2)
        table_tc.style = 'Light Shading Accent 1'
        hdr_cells_tc = table_tc.rows[0].cells
        hdr_cells_tc[0].text = 'Trouble Code'
        hdr_cells_tc[1].text = 'Claims Count'
        for tc, count in data['trouble_codes']:
            row_cells = table_tc.add_row().cells
            row_cells[0].text = str(tc)
            row_cells[1].text = str(count)
            
        # Featured Quality Cases
        h5 = doc.add_heading(level=1)
        r5 = h5.add_run('Sample Quality Cases & Technical Summaries')
        r5.font.name = 'Arial'
        r5.font.color.rgb = RGBColor(30, 58, 138)
        
        for r in data['cases']:
            ftir_no, subject, quality, dealer, summary = r
            p_case = doc.add_paragraph()
            p_case.add_run(f"FTIR No: {ftir_no}\n").bold = True
            p_case.add_run(f"Dealer: {dealer} | Quality rating: {quality}\n")
            p_case.add_run(f"Subject: {subject}\n").italic = True
            p_case.add_run(f"Summary: {summary}\n")
            
        if not data['cases']:
            doc.add_paragraph("No detailed cases available for this month.")
            
        doc.save(output_path)
        print(f"DOCX report generated at: {output_path}")



def html_escape(text):
    """Escapes HTML special characters for ReportLab Paragraph compatibility."""
    return text.replace('&', '&amp;').replace('<', '&gt;').replace('>', '&gt;')

def table_hdr_style_from_pdf():
    styles = getSampleStyleSheet()
    return ParagraphStyle(
        'TblHdr',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        textColor=colors.white,
        leading=11
    )

def table_text_style_from_pdf():
    styles = getSampleStyleSheet()
    return ParagraphStyle(
        'TblTxt',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        textColor=colors.HexColor('#334155'),
        leading=11
    )
